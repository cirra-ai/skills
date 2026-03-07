#!/usr/bin/env python3
"""
Audit Report Generator

Takes scored JSON input files produced by the AI agent and generates
deterministic HTML, DOCX, XLSX, and JSON audit reports.

Usage:
    python generate_reports.py --input-dir audit_output --output-dir audit_output \
        --org-name "Acme Corp" --org-id 00D000000000001 --instance CS42

Dependencies:
    pip install openpyxl python-docx
"""

import argparse
import html
import json
import sys
from datetime import date
from pathlib import Path

# Optional deps — fail gracefully with clear message
try:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── Brand tokens (from references/report-template.md) ─────────────────────

BRAND_BLUE = "#417AE4"
BRAND_CYAN = "#14DDDD"
BODY_BG = "#F4F6F9"
BORDER = "#E0E4EA"
MUTED = "#6B7280"

SCORE_THRESHOLDS = [
    (80, "Excellent", "#E8FBF9", "#14DDDD"),
    (70, "Good", "#E9F7EF", "#27AE60"),
    (60, "Acceptable", "#FEF3CD", "#F39C12"),
    (40, "Needs Improvement", "#FEF9E7", "#E67E22"),
    (0, "Critical", "#FDE8E8", "#E74C3C"),
]

SEVERITY_COLORS = {
    "CRITICAL": ("#FDE8E8", "#E74C3C"),
    "HIGH": ("#FEF3CD", "#E67E22"),
    "MEDIUM": ("#EBF1FB", "#417AE4"),
    "LOW": ("#E9F7EF", "#27AE60"),
}


def score_rating(pct):
    """Return (rating_label, badge_bg, badge_text) for a percentage score."""
    for threshold, label, bg, text in SCORE_THRESHOLDS:
        if pct >= threshold:
            return label, bg, text
    return SCORE_THRESHOLDS[-1][1:]


def _esc(value):
    """HTML-escape a value, converting None to empty string."""
    return html.escape(str(value)) if value is not None else ""


# ── Input loading ───────────────────────────────────────────────────────────


INPUT_FILES = {
    "counts": "counts.json",
    "apex_scores": "apex_scores.json",
    "trigger_findings": "trigger_findings.json",
    "flow_scores": "flow_scores.json",
    "process_builders": "process_builders.json",
    "lwc_scores": "lwc_scores.json",
    "permission_findings": "permission_findings.json",
    "metadata_scores": "metadata_scores.json",
    "validation_rules": "validation_rules.json",
    "workflow_rules": "workflow_rules.json",
}


def load_inputs(input_dir):
    """Load all JSON input files. Missing files become empty dicts/lists."""
    data = {}
    input_path = Path(input_dir)
    for key, filename in INPUT_FILES.items():
        fpath = input_path / filename
        if fpath.exists():
            with open(fpath) as f:
                data[key] = json.load(f)
        else:
            data[key] = {} if key == "counts" else []
    return data


# ── Score computation ───────────────────────────────────────────────────────


def compute_domain_score(items, max_score_key="max_score", score_key="score"):
    """Compute average percentage for a list of scored items."""
    if not items:
        return 0.0
    total_pct = 0.0
    for item in items:
        max_s = item.get(max_score_key, 100)
        s = item.get(score_key, 0)
        total_pct += (s / max_s * 100) if max_s > 0 else 0
    return round(total_pct / len(items), 1)


def compute_summary(data):
    """Compute the full audit summary from loaded data."""
    counts = data["counts"]

    domain_scores = {}
    for domain, items_key in [
        ("apex", "apex_scores"),
        ("flows", "flow_scores"),
        ("lwc", "lwc_scores"),
        ("metadata", "metadata_scores"),
    ]:
        items = data.get(items_key, [])
        domain_scores[domain] = compute_domain_score(
            items, max_score_key="max_score", score_key="score"
        )

    # Overall = average of domain percentages (only domains with data)
    active_scores = []
    for domain, items_key in [
        ("apex", "apex_scores"),
        ("flows", "flow_scores"),
        ("lwc", "lwc_scores"),
        ("metadata", "metadata_scores"),
    ]:
        if data.get(items_key):
            active_scores.append(domain_scores[domain])

    overall = round(sum(active_scores) / len(active_scores), 1) if active_scores else 0.0

    # Below-threshold counts (<70%)
    below_threshold = {}
    for domain, items_key in [
        ("apex", "apex_scores"),
        ("flows", "flow_scores"),
        ("lwc", "lwc_scores"),
        ("metadata", "metadata_scores"),
    ]:
        items = data.get(items_key, [])
        count = 0
        for item in items:
            max_s = item.get("max_score", 100)
            s = item.get("score", 0)
            if max_s > 0 and (s / max_s * 100) < 70:
                count += 1
        below_threshold[domain] = count

    # Severity rollup from permission findings
    severity_counts = {"CRITICAL": 0, "HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for finding in data.get("permission_findings", []):
        sev = finding.get("severity", "LOW").upper()
        if sev in severity_counts:
            severity_counts[sev] += 1

    # Top issues by domain
    top_issues = {}
    for domain, items_key in [
        ("apex", "apex_scores"),
        ("flows", "flow_scores"),
        ("lwc", "lwc_scores"),
        ("metadata", "metadata_scores"),
    ]:
        issue_counts = {}
        for item in data.get(items_key, []):
            for issue in item.get("issues", []):
                label = issue if isinstance(issue, str) else issue.get("message", str(issue))
                issue_counts[label] = issue_counts.get(label, 0) + 1
        sorted_issues = sorted(issue_counts.items(), key=lambda x: -x[1])
        top_issues[domain] = sorted_issues[:3]

    return {
        "org_name": counts.get("org_name", ""),
        "org_id": counts.get("org_id", ""),
        "instance": counts.get("instance", ""),
        "overall_score": overall,
        "overall_rating": score_rating(overall)[0],
        "domain_scores": domain_scores,
        "counts": counts,
        "below_threshold": below_threshold,
        "severity_counts": severity_counts,
        "top_issues": {k: [{"issue": i, "count": c} for i, c in v] for k, v in top_issues.items()},
    }


# ── HTML report ─────────────────────────────────────────────────────────────


def _severity_badge_html(severity):
    bg, fg = SEVERITY_COLORS.get(severity.upper(), ("#EBF1FB", "#417AE4"))
    return (
        f'<span style="background:{bg};color:{fg};padding:2px 8px;'
        f'border-radius:10px;font-size:11px;font-weight:600">{_esc(severity)}</span>'
    )


def _score_badge_html(score, max_score):
    pct = (score / max_score * 100) if max_score > 0 else 0
    _, bg, fg = score_rating(pct)
    return (
        f'<span style="background:{bg};color:{fg};padding:2px 8px;'
        f'border-radius:10px;font-size:11px;font-weight:600">'
        f"{score}/{max_score}</span>"
    )


def _table_html(headers, rows):
    """Build an HTML table with branded headers."""
    parts = ["<table>", "<thead><tr>"]
    for h in headers:
        parts.append(f"<th>{_esc(h)}</th>")
    parts.append("</tr></thead><tbody>")
    for row in rows:
        parts.append("<tr>")
        for cell in row:
            parts.append(f"<td>{cell}</td>")  # cell may contain HTML badges
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "\n".join(parts)


def _findings_html(findings):
    """Render a list of findings with severity left-border strips."""
    if not findings:
        return "<p>No findings.</p>"
    parts = []
    for f in findings:
        sev = f.get("severity", "MEDIUM").lower()
        css_class = {
            "critical": "critical",
            "high": "warning",
            "medium": "info",
            "low": "positive",
        }.get(sev, "info")
        parts.append(
            f'<div class="finding {css_class}">'
            f'<span class="finding-badge">{_esc(f.get("severity", "MEDIUM"))}</span> '
            f'{_esc(f.get("message", f.get("finding", "")))}'
            f"</div>"
        )
    return "\n".join(parts)


def _recommendations_html(summary, data):
    """Build the top-10 recommendations section."""
    recs = []

    # Collect from below-threshold items
    for domain, items_key, default_max in [
        ("Apex", "apex_scores", 150),
        ("Flows", "flow_scores", 110),
        ("LWC", "lwc_scores", 165),
        ("Metadata", "metadata_scores", 120),
    ]:
        for item in sorted(data.get(items_key, []), key=lambda x: x.get("score", 0)):
            s = item.get("score", 0)
            max_s = item.get("max_score", default_max)
            pct = (s / max_s * 100) if max_s > 0 else 0
            if pct < 70:
                name = item.get("name", "Unknown")
                top_issue = ""
                issues = item.get("issues", [])
                if issues:
                    top_issue = issues[0] if isinstance(issues[0], str) else issues[0].get("message", "")
                recs.append(f"[{domain}] Fix **{_esc(name)}** (score {s}/{max_s}). {_esc(top_issue)}")

    # Add permission findings (CRITICAL first)
    for f in sorted(
        data.get("permission_findings", []),
        key=lambda x: {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}.get(
            x.get("severity", "LOW").upper(), 4
        ),
    ):
        recs.append(f"[Permissions] {_esc(f.get('message', f.get('finding', '')))}")

    # Legacy automation
    pb_count = len(data.get("process_builders", []))
    wr_count = len(data.get("workflow_rules", []))
    if pb_count:
        recs.append(f"[Automation] Migrate {pb_count} active Process Builder(s) to Flow")
    if wr_count:
        recs.append(f"[Automation] Migrate {wr_count} Workflow Rule(s) to Flow")

    recs = recs[:10]
    if not recs:
        return "<p>No recommendations — org is in great shape.</p>"

    parts = []
    for i, rec in enumerate(recs, 1):
        parts.append(
            f'<div style="display:flex;gap:12px;align-items:flex-start;margin-bottom:10px">'
            f'<div class="rec-num">{i}</div>'
            f"<div>{rec}</div></div>"
        )
    return "\n".join(parts)


def generate_html(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the self-contained HTML audit report."""
    overall = summary["overall_score"]
    rating = summary["overall_rating"]
    counts = summary["counts"]

    # Core CSS is embedded below (from references/report-template.md)

    sections = []

    # ── Executive Summary ──
    inv_lines = []
    inv_map = [
        ("apex_classes", "Apex Classes"),
        ("apex_triggers", "Apex Triggers"),
        ("active_flows", "Active Flows"),
        ("process_builders", "Process Builders"),
        ("lwc_bundles", "LWC Components"),
        ("custom_objects", "Custom Objects"),
        ("validation_rules", "Validation Rules"),
        ("workflow_rules", "Workflow Rules"),
        ("permission_sets", "Permission Sets"),
        ("permission_set_groups", "Permission Set Groups"),
        ("profiles", "Profiles"),
    ]
    for key, label in inv_map:
        val = counts.get(key, 0)
        inv_lines.append(f"<li><strong>{val}</strong> {label}</li>")
    sections.append(
        f'<div class="card"><h2>Executive Summary</h2>'
        f'<ul>{"".join(inv_lines)}</ul></div>'
    )

    # ── Domain score cards ──
    domain_cards = []
    for domain, label, _max_s in [
        ("apex", "Apex", 150),
        ("flows", "Flows", 110),
        ("lwc", "LWC", 165),
        ("metadata", "Metadata", 120),
    ]:
        pct = summary["domain_scores"].get(domain, 0)
        r, bg, fg = score_rating(pct)
        below = summary["below_threshold"].get(domain, 0)
        domain_cards.append(
            f'<div class="score-card">'
            f'<div class="score-card-label">{label}</div>'
            f'<div style="font-size:28px;font-weight:700;color:{fg}">{pct}%</div>'
            f'<div style="font-size:11px;color:{MUTED}">{r}</div>'
            f'<div style="font-size:11px;color:{MUTED}">{below} below 70%</div>'
            f"</div>"
        )
    sections.append(
        f'<div class="card"><h2>Domain Scores</h2>'
        f'<div style="display:flex;gap:16px;flex-wrap:wrap">{"".join(domain_cards)}</div></div>'
    )

    # ── Apex Classes ──
    if data.get("apex_scores"):
        rows = []
        for item in sorted(data["apex_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 150)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Apex Classes</h2>'
            f'{_table_html(["Name", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Apex Triggers ──
    if data.get("trigger_findings"):
        rows = []
        for item in data["trigger_findings"]:
            findings_str = "; ".join(
                f.get("message", f.get("finding", ""))
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            )
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(item.get("events", "")),
                _severity_badge_html(sev),
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Apex Triggers</h2>'
            f'{_table_html(["Name", "Object", "Events", "Severity", "Findings"], rows)}</div>'
        )

    # ── Flows ──
    if data.get("flow_scores"):
        rows = []
        for item in sorted(data["flow_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("process_type", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 110)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Flows</h2>'
            f'{_table_html(["Name", "Process Type", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Process Builders ──
    if data.get("process_builders"):
        rows = []
        for item in data["process_builders"]:
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(str(item.get("criteria_count", ""))),
                _esc(item.get("actions_summary", "")),
                _severity_badge_html(item.get("migration_priority", "HIGH")),
            ])
        sections.append(
            f'<div class="card"><h2>Process Builders</h2>'
            f'{_table_html(["Name", "Object", "Criteria", "Actions", "Migration Priority"], rows)}</div>'
        )

    # ── LWC ──
    if data.get("lwc_scores"):
        rows = []
        for item in sorted(data["lwc_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 165)),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Lightning Web Components</h2>'
            f'{_table_html(["Name", "Score", "Top Issues"], rows)}</div>'
        )

    # ── Permissions ──
    if data.get("permission_findings"):
        sections.append(
            f'<div class="card"><h2>Profiles &amp; Permissions</h2>'
            f"{_findings_html(data['permission_findings'])}</div>"
        )

    # ── Metadata / Data Model ──
    if data.get("metadata_scores"):
        rows = []
        for item in sorted(data["metadata_scores"], key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            rows.append([
                _esc(item.get("name", "")),
                _score_badge_html(item.get("score", 0), item.get("max_score", 120)),
                _esc(str(item.get("field_count", ""))),
                _esc(str(item.get("relationship_count", ""))),
                _esc(issues_str),
            ])
        sections.append(
            f'<div class="card"><h2>Data Model</h2>'
            f'{_table_html(["Object", "Score", "Fields", "Relationships", "Top Issues"], rows)}</div>'
        )

    # ── Validation Rules ──
    if data.get("validation_rules"):
        rows = []
        for item in data["validation_rules"]:
            findings_str = "; ".join(
                f.get("message", f.get("finding", ""))
                for f in item.get("findings", [])[:3]
            )
            sev = max(
                (f.get("severity", "LOW") for f in item.get("findings", [])),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
                default="LOW",
            ) if item.get("findings") else "LOW"
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(str(item.get("active", ""))),
                _severity_badge_html(sev),
                _esc(findings_str),
            ])
        sections.append(
            f'<div class="card"><h2>Validation Rules</h2>'
            f'{_table_html(["Name", "Object", "Active", "Severity", "Findings"], rows)}</div>'
        )

    # ── Workflow Rules ──
    if data.get("workflow_rules"):
        rows = []
        for item in data["workflow_rules"]:
            rows.append([
                _esc(item.get("name", "")),
                _esc(item.get("object", "")),
                _esc(item.get("action_types", "")),
                _severity_badge_html(item.get("migration_priority", "HIGH")),
            ])
        sections.append(
            f'<div class="card"><h2>Workflow Rules</h2>'
            f'{_table_html(["Name", "Object", "Action Types", "Migration Priority"], rows)}</div>'
        )

    # ── Recommendations ──
    sections.append(
        f'<div class="card"><h2>Recommendations</h2>'
        f"{_recommendations_html(summary, data)}</div>"
    )

    body = "\n".join(sections)

    report = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Salesforce Org Audit Report — {_esc(org_name)}</title>
<style>
:root {{
  --brand-blue: {BRAND_BLUE};
  --brand-cyan: {BRAND_CYAN};
}}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       background: {BODY_BG}; color: #1a1a2e; line-height: 1.5; }}

.banner {{
  background: {BRAND_BLUE};
  padding: 28px 40px 32px;
  display: flex; align-items: center; gap: 24px;
}}
.banner-text {{ flex: 1; }}
.banner-title {{ font-size: 26px; font-weight: 700; color: #fff; letter-spacing: -0.3px; }}
.banner-subtitle {{ font-size: 13px; color: rgba(255,255,255,.80); margin-top: 4px; }}
.banner-score {{ text-align: right; flex-shrink: 0; }}
.banner-score-value {{ font-size: 48px; font-weight: 800; color: #fff; line-height: 1; }}
.banner-score-label {{ font-size: 11px; text-transform: uppercase; letter-spacing: 1px;
                       color: rgba(255,255,255,.75); margin-top: 4px; }}
.banner-score-rating {{
  display: inline-block; margin-top: 6px;
  background: rgba(255,255,255,.20); border: 1px solid rgba(255,255,255,.35);
  border-radius: 20px; padding: 2px 12px;
  font-size: 12px; font-weight: 600; color: #fff;
}}

.card {{ background: #fff; border: 1px solid {BORDER}; border-radius: 8px;
         padding: 24px; margin: 16px 40px; }}
.card h2 {{ color: {BRAND_BLUE}; font-size: 18px; margin-bottom: 16px; }}

table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
thead th {{
  background: {BRAND_BLUE}; color: #fff; font-size: 11px; font-weight: 700;
  text-transform: uppercase; letter-spacing: 0.6px;
  padding: 10px 14px; text-align: left;
}}
tbody td {{ padding: 10px 14px; border-bottom: 1px solid {BORDER}; font-size: 13px; }}
tbody tr:hover {{ background: #f8f9fb; }}

.score-card {{
  background: #fff; border: 1px solid {BORDER}; border-radius: 8px;
  padding: 16px 20px; min-width: 140px; text-align: center;
}}
.score-card-label {{ font-size: 12px; font-weight: 600; color: {MUTED};
                     text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 4px; }}

.rec-num {{
  width: 28px; height: 28px; border-radius: 50%; background: {BRAND_BLUE};
  color: #fff; font-size: 13px; font-weight: 700;
  display: flex; align-items: center; justify-content: center; flex-shrink: 0;
}}

.finding {{ border-left: 4px solid; border-radius: 6px; padding: 12px 16px; margin-bottom: 8px;
            background: #fff; }}
.finding.critical {{ border-left-color: #E74C3C; }}
.finding.warning {{ border-left-color: #E67E22; }}
.finding.info {{ border-left-color: {BRAND_BLUE}; }}
.finding.positive {{ border-left-color: #27AE60; }}
.finding-badge {{
  display: inline-block; padding: 2px 8px; border-radius: 10px;
  font-size: 11px; font-weight: 600; margin-right: 8px;
}}
.finding.critical .finding-badge {{ background: #FDE8E8; color: #E74C3C; }}
.finding.warning .finding-badge {{ background: #FEF3CD; color: #E67E22; }}
.finding.info .finding-badge {{ background: #EBF1FB; color: {BRAND_BLUE}; }}
.finding.positive .finding-badge {{ background: #E9F7EF; color: #27AE60; }}

.footer {{
  background: {BRAND_BLUE};
  padding: 20px 40px; text-align: center;
  font-size: 12px; color: rgba(255,255,255,.80); margin-top: 24px;
}}
.footer a {{ color: #fff; }}
</style>
</head>
<body>
<div class="banner">
  <div class="banner-text">
    <div class="banner-title">Salesforce Org Audit Report</div>
    <div class="banner-subtitle">
      {_esc(org_name)} &middot; Org ID: {_esc(org_id)} &middot;
      Instance: {_esc(instance)} &middot; {_esc(run_date)}
    </div>
  </div>
  <div class="banner-score">
    <div class="banner-score-value">{overall:.0f}</div>
    <div class="banner-score-label">out of 100</div>
    <div class="banner-score-rating">{_esc(rating)}</div>
  </div>
</div>
{body}
<div class="footer">
  Generated by <a href="https://cirra.ai">Cirra AI</a> Audit Engine
  &nbsp;&middot;&nbsp; {_esc(run_date)} &nbsp;&middot;&nbsp;
  Org: {_esc(org_name)} ({_esc(org_id)})
</div>
</body>
</html>"""

    Path(output_path).write_text(report, encoding="utf-8")
    return output_path


# ── DOCX report ─────────────────────────────────────────────────────────────


def _hex_to_rgb(hex_color):
    """Convert '#RRGGBB' or 'RRGGBB' to RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def generate_docx(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the Word audit report."""
    if not HAS_DOCX:
        print("WARNING: python-docx not installed — skipping DOCX generation", file=sys.stderr)
        return None

    doc = docx.Document()

    # Page setup — US Letter
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = doc.add_heading("Salesforce Org Audit Report", level=0)
    for run in title.runs:
        run.font.color.rgb = _hex_to_rgb(BRAND_BLUE)

    doc.add_paragraph(
        f"{org_name} | Org ID: {org_id} | Instance: {instance} | {run_date}"
    )

    overall = summary["overall_score"]
    rating = summary["overall_rating"]
    p = doc.add_paragraph()
    run = p.add_run(f"Overall Score: {overall:.0f}/100 — {rating}")
    run.bold = True
    run.font.size = Pt(14)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    counts = summary["counts"]
    inv_map = [
        ("apex_classes", "Apex Classes"),
        ("apex_triggers", "Apex Triggers"),
        ("active_flows", "Active Flows"),
        ("process_builders", "Process Builders"),
        ("lwc_bundles", "LWC Components"),
        ("custom_objects", "Custom Objects"),
        ("validation_rules", "Validation Rules"),
        ("workflow_rules", "Workflow Rules"),
        ("permission_sets", "Permission Sets"),
        ("permission_set_groups", "Permission Set Groups"),
        ("profiles", "Profiles"),
    ]
    for key, label in inv_map:
        doc.add_paragraph(f"{counts.get(key, 0)} {label}", style="List Bullet")

    # Domain Scores
    doc.add_heading("Domain Scores", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Domain", "Score %", "Rating", "Below 70%"]):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True

    for domain, label in [("apex", "Apex"), ("flows", "Flows"), ("lwc", "LWC"), ("metadata", "Metadata")]:
        pct = summary["domain_scores"].get(domain, 0)
        r = score_rating(pct)[0]
        below = summary["below_threshold"].get(domain, 0)
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{pct}%"
        row[2].text = r
        row[3].text = str(below)

    # Scored sections
    for section_title, items_key, max_label in [
        ("Apex Classes", "apex_scores", 150),
        ("Flows", "flow_scores", 110),
        ("LWC Components", "lwc_scores", 165),
        ("Data Model", "metadata_scores", 120),
    ]:
        items = data.get(items_key, [])
        if not items:
            continue
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(["Name", "Score", "Top Issues"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for item in sorted(items, key=lambda x: x.get("score", 0)):
            issues_str = "; ".join(
                i if isinstance(i, str) else i.get("message", "")
                for i in item.get("issues", [])[:3]
            )
            row = table.add_row().cells
            row[0].text = item.get("name", "")
            row[1].text = f"{item.get('score', 0)}/{item.get('max_score', max_label)}"
            row[2].text = issues_str

    # Permission Findings
    if data.get("permission_findings"):
        doc.add_heading("Profiles & Permissions", level=1)
        for f in data["permission_findings"]:
            doc.add_paragraph(
                f"[{f.get('severity', 'MEDIUM')}] {f.get('message', f.get('finding', ''))}",
                style="List Bullet",
            )

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Generated by Cirra AI Audit Engine | {run_date} | Org: {org_name} ({org_id})")
    run.font.size = Pt(9)
    run.font.color.rgb = _hex_to_rgb(MUTED)

    doc.save(output_path)
    return output_path


# ── XLSX report ─────────────────────────────────────────────────────────────


def generate_xlsx(data, summary, org_name, org_id, instance, run_date, output_path):
    """Generate the Excel audit workbook."""
    if not HAS_OPENPYXL:
        print("WARNING: openpyxl not installed — skipping XLSX generation", file=sys.stderr)
        return None

    wb = openpyxl.Workbook()

    header_fill = PatternFill("solid", fgColor="417AE4")
    header_font = Font(bold=True, color="FFFFFF", name="Arial", size=11)

    def apply_header(ws, headers):
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="left", vertical="center")

    def auto_width(ws):
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    score_fills = {
        "Excellent": PatternFill("solid", fgColor="E8FBF9"),
        "Good": PatternFill("solid", fgColor="E9F7EF"),
        "Acceptable": PatternFill("solid", fgColor="FEF3CD"),
        "Warning": PatternFill("solid", fgColor="FEF9E7"),
        "Critical": PatternFill("solid", fgColor="FDE8E8"),
    }
    score_fonts = {
        "Excellent": Font(color="14DDDD", bold=True, name="Arial", size=11),
        "Good": Font(color="27AE60", bold=True, name="Arial", size=11),
        "Acceptable": Font(color="F39C12", bold=True, name="Arial", size=11),
        "Warning": Font(color="E67E22", bold=True, name="Arial", size=11),
        "Critical": Font(color="E74C3C", bold=True, name="Arial", size=11),
    }

    def style_score_cell(cell, score, max_score):
        pct = (score / max_score * 100) if max_score > 0 else 0
        rating = score_rating(pct)[0]
        # Map "Needs Improvement" to "Warning" for styling
        style_key = rating if rating in score_fills else "Warning"
        cell.fill = score_fills.get(style_key, score_fills["Warning"])
        cell.font = score_fonts.get(style_key, score_fonts["Warning"])

    # Sheet 1 — Apex Classes
    ws = wb.active
    ws.title = "Apex Classes"
    ws.sheet_properties.tabColor = "417AE4"
    apply_header(ws, ["Name", "Score", "Max", "Rating", "Top Issues"])
    for i, item in enumerate(sorted(data.get("apex_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 150)
        pct = (s / m * 100) if m > 0 else 0
        ws.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws.cell(row=i, column=3, value=m)
        ws.cell(row=i, column=4, value=score_rating(pct)[0])
        issues = item.get("issues", [])
        ws.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in issues[:3]
        ))
    auto_width(ws)

    # Sheet 2 — Apex Triggers
    ws2 = wb.create_sheet("Apex Triggers")
    ws2.sheet_properties.tabColor = "417AE4"
    apply_header(ws2, ["Name", "Object", "Events", "Findings", "Severity"])
    for i, item in enumerate(data.get("trigger_findings", []), 2):
        ws2.cell(row=i, column=1, value=item.get("name", ""))
        ws2.cell(row=i, column=2, value=item.get("object", ""))
        ws2.cell(row=i, column=3, value=item.get("events", ""))
        findings = item.get("findings", [])
        ws2.cell(row=i, column=4, value="; ".join(
            f.get("message", f.get("finding", "")) for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws2.cell(row=i, column=5, value=sev)
    auto_width(ws2)

    # Sheet 3 — Flows
    ws3 = wb.create_sheet("Flows")
    ws3.sheet_properties.tabColor = "417AE4"
    apply_header(ws3, ["Name", "Process Type", "Score", "Max", "Top Issues"])
    for i, item in enumerate(sorted(data.get("flow_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 110)
        ws3.cell(row=i, column=1, value=item.get("name", ""))
        ws3.cell(row=i, column=2, value=item.get("process_type", ""))
        sc = ws3.cell(row=i, column=3, value=s)
        style_score_cell(sc, s, m)
        ws3.cell(row=i, column=4, value=m)
        ws3.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws3)

    # Sheet 4 — Process Builders
    ws4 = wb.create_sheet("Process Builders")
    ws4.sheet_properties.tabColor = "417AE4"
    apply_header(ws4, ["Name", "Object", "Criteria Count", "Actions", "Migration Priority"])
    for i, item in enumerate(data.get("process_builders", []), 2):
        ws4.cell(row=i, column=1, value=item.get("name", ""))
        ws4.cell(row=i, column=2, value=item.get("object", ""))
        ws4.cell(row=i, column=3, value=item.get("criteria_count", 0))
        ws4.cell(row=i, column=4, value=item.get("actions_summary", ""))
        ws4.cell(row=i, column=5, value=item.get("migration_priority", "HIGH"))
    auto_width(ws4)

    # Sheet 5 — LWC
    ws5 = wb.create_sheet("LWC")
    ws5.sheet_properties.tabColor = "417AE4"
    apply_header(ws5, ["Name", "Score", "Max", "Rating", "Top Issues"])
    for i, item in enumerate(sorted(data.get("lwc_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 165)
        pct = (s / m * 100) if m > 0 else 0
        ws5.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws5.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws5.cell(row=i, column=3, value=m)
        ws5.cell(row=i, column=4, value=score_rating(pct)[0])
        ws5.cell(row=i, column=5, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws5)

    # Sheet 6 — Profiles
    ws6 = wb.create_sheet("Profiles")
    ws6.sheet_properties.tabColor = "417AE4"
    apply_header(ws6, ["Name", "UserType", "Key Permissions", "Findings", "Severity"])
    # Profiles come from permission_findings filtered
    profile_findings = [f for f in data.get("permission_findings", []) if f.get("type") == "profile"]
    for i, item in enumerate(profile_findings, 2):
        ws6.cell(row=i, column=1, value=item.get("name", ""))
        ws6.cell(row=i, column=2, value=item.get("user_type", ""))
        ws6.cell(row=i, column=3, value=item.get("key_permissions", ""))
        ws6.cell(row=i, column=4, value=item.get("message", item.get("finding", "")))
        ws6.cell(row=i, column=5, value=item.get("severity", ""))
    auto_width(ws6)

    # Sheet 7 — Permission Sets
    ws7 = wb.create_sheet("Permission Sets")
    ws7.sheet_properties.tabColor = "417AE4"
    apply_header(ws7, ["Name", "Label", "Assignments", "Findings", "Severity"])
    ps_findings = [f for f in data.get("permission_findings", []) if f.get("type") == "permission_set"]
    for i, item in enumerate(ps_findings, 2):
        ws7.cell(row=i, column=1, value=item.get("name", ""))
        ws7.cell(row=i, column=2, value=item.get("label", ""))
        ws7.cell(row=i, column=3, value=item.get("assignments", 0))
        ws7.cell(row=i, column=4, value=item.get("message", item.get("finding", "")))
        ws7.cell(row=i, column=5, value=item.get("severity", ""))
    auto_width(ws7)

    # Sheet 8 — Custom Objects
    ws8 = wb.create_sheet("Custom Objects")
    ws8.sheet_properties.tabColor = "417AE4"
    apply_header(ws8, ["Name", "Score", "Max", "Fields", "Relationships", "Top Issues"])
    for i, item in enumerate(sorted(data.get("metadata_scores", []), key=lambda x: x.get("score", 0)), 2):
        s = item.get("score", 0)
        m = item.get("max_score", 120)
        ws8.cell(row=i, column=1, value=item.get("name", ""))
        sc = ws8.cell(row=i, column=2, value=s)
        style_score_cell(sc, s, m)
        ws8.cell(row=i, column=3, value=m)
        ws8.cell(row=i, column=4, value=item.get("field_count", 0))
        ws8.cell(row=i, column=5, value=item.get("relationship_count", 0))
        ws8.cell(row=i, column=6, value="; ".join(
            x if isinstance(x, str) else x.get("message", "") for x in item.get("issues", [])[:3]
        ))
    auto_width(ws8)

    # Sheet 9 — Validation Rules
    ws9 = wb.create_sheet("Validation Rules")
    ws9.sheet_properties.tabColor = "417AE4"
    apply_header(ws9, ["Name", "Object", "Active", "Findings", "Severity"])
    for i, item in enumerate(data.get("validation_rules", []), 2):
        ws9.cell(row=i, column=1, value=item.get("name", ""))
        ws9.cell(row=i, column=2, value=item.get("object", ""))
        ws9.cell(row=i, column=3, value=str(item.get("active", "")))
        findings = item.get("findings", [])
        ws9.cell(row=i, column=4, value="; ".join(
            f.get("message", f.get("finding", "")) for f in findings[:3]
        ))
        if findings:
            sev = max(
                (f.get("severity", "LOW") for f in findings),
                key=lambda s: {"CRITICAL": 4, "HIGH": 3, "MEDIUM": 2, "LOW": 1}.get(s.upper(), 0),
            )
            ws9.cell(row=i, column=5, value=sev)
    auto_width(ws9)

    # Sheet 10 — Workflow Rules
    ws10 = wb.create_sheet("Workflow Rules")
    ws10.sheet_properties.tabColor = "417AE4"
    apply_header(ws10, ["Name", "Object", "Action Types", "Migration Priority"])
    for i, item in enumerate(data.get("workflow_rules", []), 2):
        ws10.cell(row=i, column=1, value=item.get("name", ""))
        ws10.cell(row=i, column=2, value=item.get("object", ""))
        ws10.cell(row=i, column=3, value=item.get("action_types", ""))
        ws10.cell(row=i, column=4, value=item.get("migration_priority", "HIGH"))
    auto_width(ws10)

    # Sheet 11 — Summary
    ws11 = wb.create_sheet("Summary")
    ws11.sheet_properties.tabColor = "417AE4"
    apply_header(ws11, ["Metric", "Value"])
    summary_rows = [
        ("Org Name", org_name),
        ("Org ID", org_id),
        ("Instance", instance),
        ("Run Date", run_date),
        ("Overall Score", f"{summary['overall_score']:.0f}/100"),
        ("Overall Rating", summary["overall_rating"]),
        ("", ""),
        ("Domain Scores", ""),
    ]
    for domain in ["apex", "flows", "lwc", "metadata"]:
        pct = summary["domain_scores"].get(domain, 0)
        summary_rows.append((f"  {domain.title()}", f"{pct}%"))

    summary_rows.append(("", ""))
    summary_rows.append(("Severity Counts", ""))
    for sev in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]:
        summary_rows.append((f"  {sev}", summary["severity_counts"].get(sev, 0)))

    for i, (metric, value) in enumerate(summary_rows, 2):
        ws11.cell(row=i, column=1, value=metric)
        ws11.cell(row=i, column=2, value=value)
    auto_width(ws11)

    wb.save(output_path)
    return output_path


# ── JSON summary ────────────────────────────────────────────────────────────


def generate_json_summary(summary, run_date, output_path):
    """Generate the JSON audit summary."""
    output = {**summary, "generated_date": run_date}
    Path(output_path).write_text(json.dumps(output, indent=2, default=str), encoding="utf-8")
    return output_path


# ── Main ────────────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(description="Generate Salesforce Org Audit Reports")
    parser.add_argument("--input-dir", required=True, help="Directory containing scored JSON files")
    parser.add_argument("--output-dir", required=True, help="Directory for generated reports")
    parser.add_argument("--org-name", required=True, help="Salesforce org name")
    parser.add_argument("--org-id", default="", help="18-char Org ID")
    parser.add_argument("--instance", default="", help="Salesforce instance (e.g. CS42)")
    parser.add_argument("--run-date", default=None, help="Audit date (YYYY-MM-DD). Defaults to today.")
    args = parser.parse_args()

    run_date = args.run_date or date.today().isoformat()

    # Validate input dir
    input_path = Path(args.input_dir)
    if not input_path.is_dir():
        print(f"ERROR: Input directory does not exist: {args.input_dir}", file=sys.stderr)
        sys.exit(1)

    # Create output dir
    output_path = Path(args.output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    # Load and compute
    data = load_inputs(args.input_dir)
    summary = compute_summary(data)

    # Override org info from CLI args (takes precedence over counts.json)
    org_name = args.org_name
    org_id = args.org_id or summary.get("org_id", "")
    instance = args.instance or summary.get("instance", "")

    generated = []

    # HTML
    html_path = output_path / "Salesforce_Org_Audit_Report.html"
    generate_html(data, summary, org_name, org_id, instance, run_date, html_path)
    generated.append(str(html_path))
    print(f"  HTML:  {html_path}")

    # DOCX
    docx_path = output_path / "Salesforce_Org_Audit_Report.docx"
    result = generate_docx(data, summary, org_name, org_id, instance, run_date, docx_path)
    if result:
        generated.append(str(docx_path))
        print(f"  DOCX:  {docx_path}")

    # XLSX
    xlsx_path = output_path / "Salesforce_Org_Audit_Scores.xlsx"
    result = generate_xlsx(data, summary, org_name, org_id, instance, run_date, xlsx_path)
    if result:
        generated.append(str(xlsx_path))
        print(f"  XLSX:  {xlsx_path}")

    # JSON summary
    json_path = output_path / "audit_summary.json"
    generate_json_summary(summary, run_date, json_path)
    generated.append(str(json_path))
    print(f"  JSON:  {json_path}")

    print(f"\nOverall Score: {summary['overall_score']:.0f}/100 — {summary['overall_rating']}")
    print(f"Generated {len(generated)} report(s)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
